import io
import json
import random
import re
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse
from urllib.request import urlopen

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE_DIR = Path(__file__).parent
COLUNAS_OBRIGATORIAS = {
    "id", "disciplina", "assunto", "texto",
    "A", "B", "C", "D", "E", "gabarito", "dificuldade", "modelos",
}


def carregar_config() -> dict:
    with open(BASE_DIR / "config_modelos.json", encoding="utf-8") as f:
        return json.load(f)


def carregar_banco(caminho_csv: str | Path) -> pd.DataFrame:
    df = pd.read_csv(caminho_csv, dtype=str)
    faltando = COLUNAS_OBRIGATORIAS - set(df.columns)
    if faltando:
        raise ValueError(f"Colunas ausentes no CSV: {', '.join(sorted(faltando))}")
    df["modelos"] = df["modelos"].fillna("")
    df["E"] = df["E"].fillna("")
    df["id"] = df["id"].fillna("").astype(str).str.strip()
    if (df["id"] == "").any():
        raise ValueError("Há questões com ID vazio. Preencha a coluna 'id' em todas as linhas.")
    duplicados = df["id"].duplicated(keep=False)
    if duplicados.any():
        ids_dup = ", ".join(sorted(df.loc[duplicados, "id"].unique())[:10])
        raise ValueError(f"IDs duplicados no CSV: {ids_dup}")
    df["dificuldade"] = pd.to_numeric(df["dificuldade"], errors="coerce").astype("Int64")
    invalidas = df["dificuldade"].isna() | ~df["dificuldade"].between(1, 5)
    if invalidas.any():
        raise ValueError(
            f"Dificuldade inválida em {invalidas.sum()} questão(ões). Use inteiros de 1 a 5."
        )
    gabarito = df["gabarito"].fillna("").astype(str).str.strip().str.upper()
    invalidos_gabarito = ~gabarito.isin(["A", "B", "C", "D", "E"])
    if invalidos_gabarito.any():
        raise ValueError(
            f"Gabarito inválido em {invalidos_gabarito.sum()} questão(ões). Use apenas A, B, C, D ou E."
        )
    df["gabarito"] = gabarito
    return df


# ── distribuição ──────────────────────────────────────────────────────────────

def distribuicao_efetiva(modelo: dict) -> dict[int, float]:
    """Returns the active difficulty distribution for a model (config or uniform fallback)."""
    dif_range = modelo.get("dificuldade_range", [1, 5])
    dist_cfg = modelo.get("dificuldade_distribuicao")
    if dist_cfg:
        return {int(k): float(v) for k, v in dist_cfg.items()}
    levels = list(range(dif_range[0], dif_range[1] + 1))
    return {lv: 1.0 / len(levels) for lv in levels}


def _calcular_cotas(quantidade: int, distribuicao: dict[int, float]) -> dict[int, int]:
    """Largest-remainder method: distributes n questions across difficulty levels."""
    reais = {lv: quantidade * pct for lv, pct in distribuicao.items()}
    cotas = {lv: int(v) for lv, v in reais.items()}
    resto = quantidade - sum(cotas.values())
    por_fracao = sorted(reais.items(), key=lambda x: -(x[1] - int(x[1])))
    for i in range(resto):
        cotas[por_fracao[i % len(por_fracao)][0]] += 1
    return cotas


# ── filtragem e seleção ───────────────────────────────────────────────────────

def _filtrar(df: pd.DataFrame, disciplina: str, modelo_nome: str) -> pd.DataFrame:
    disc_lower = disciplina.strip().lower()
    mask_disc = df["disciplina"].str.strip().str.lower() == disc_lower

    def compativel(m: str) -> bool:
        m = m.strip()
        return m == "" or modelo_nome in [x.strip().lower() for x in m.split(";")]

    mask_modelo = df["modelos"].apply(compativel)
    return df[mask_disc & mask_modelo]


def _normalizar_assuntos(assuntos_cfg) -> set[str]:
    if assuntos_cfg is None:
        return set()
    if isinstance(assuntos_cfg, str):
        return {assuntos_cfg.strip().lower()} if assuntos_cfg.strip() else set()
    if isinstance(assuntos_cfg, list):
        return {
            str(a).strip().lower()
            for a in assuntos_cfg
            if str(a).strip()
        }
    return set()


def _filtrar_por_assunto(pool: pd.DataFrame, bloco: dict) -> pd.DataFrame:
    assuntos = _normalizar_assuntos(bloco.get("assuntos", bloco.get("assunto")))
    if not assuntos:
        return pool
    return pool[pool["assunto"].fillna("").str.strip().str.lower().isin(assuntos)]


def _questao_tem_imagem(row: dict) -> bool:
    campos = [str(row.get("texto", ""))]
    for letra in ("A", "B", "C", "D", "E"):
        campos.append(str(row.get(letra, "")))
    return any("[img]" in campo.lower() for campo in campos)


def _normalizar_regras_assunto(
    regras_assunto: dict[str, dict[str, int]] | None,
) -> dict[str, dict[str, int]]:
    if not regras_assunto:
        return {}
    normalizado: dict[str, dict[str, int]] = {}
    for disciplina, assuntos in regras_assunto.items():
        disc_key = str(disciplina).strip().lower()
        if not disc_key or not isinstance(assuntos, dict):
            continue
        bucket: dict[str, int] = {}
        for assunto, minimo in assuntos.items():
            assunto_key = str(assunto).strip().lower()
            try:
                minimo_int = int(minimo)
            except (TypeError, ValueError):
                continue
            if assunto_key and minimo_int > 0:
                bucket[assunto_key] = minimo_int
        if bucket:
            normalizado[disc_key] = bucket
    return normalizado


def _selecionar_questoes(
    pool: pd.DataFrame,
    disciplina: str,
    quantidade: int,
    dif_range: list[int],
    distribuicao: dict[int, float],
    rng: random.Random,
    usadas: set[str],
    regras_assunto: dict[str, dict[str, int]] | None = None,
    aplicar_distribuicao: bool = True,
    forcar_questao_imagem: bool = False,
) -> list[dict]:
    """Selects questions by difficulty distribution."""
    pool_valido = pool[
        ~pool["id"].isin(usadas) & pool["dificuldade"].between(dif_range[0], dif_range[1])
    ]
    regras_norm = _normalizar_regras_assunto(regras_assunto)
    regras_disc = regras_norm.get(disciplina.strip().lower(), {})
    selecionadas: list[dict] = []
    ids_sel: set[str] = set()

    # Primeiro, cumpre mínimos por assunto da disciplina (se houver).
    for assunto, minimo in regras_disc.items():
        if minimo <= 0:
            continue
        grupo_assunto = pool_valido[
            pool_valido["assunto"].fillna("").str.strip().str.lower().eq(assunto)
            & (~pool_valido["id"].isin(ids_sel))
        ].to_dict("records")
        escolhidos = rng.sample(grupo_assunto, min(minimo, len(grupo_assunto)))
        selecionadas.extend(escolhidos)
        ids_sel.update(str(q["id"]) for q in escolhidos)

    faltam_total = quantidade - len(selecionadas)
    if faltam_total <= 0:
        return selecionadas[:quantidade]

    if not aplicar_distribuicao:
        restantes = pool_valido[~pool_valido["id"].isin(ids_sel)].to_dict("records")
        complemento = rng.sample(restantes, min(faltam_total, len(restantes)))
        selecionadas.extend(complemento)
        return selecionadas

    cotas = _calcular_cotas(faltam_total, distribuicao)

    for nivel, cota in cotas.items():
        if cota == 0:
            continue
        grupo = pool_valido[
            (pool_valido["dificuldade"] == nivel) & (~pool_valido["id"].isin(ids_sel))
        ].to_dict("records")
        if len(grupo) < cota:
            raise ValueError(
                f"Dificuldade {nivel}: necessário {cota}, disponível {len(grupo)}. "
                "Inclua mais questões desse nível ou ajuste a distribuição."
            )
        escolhidos = rng.sample(grupo, cota)
        selecionadas.extend(escolhidos)
        ids_sel.update(str(q["id"]) for q in escolhidos)

    if forcar_questao_imagem and selecionadas and not any(_questao_tem_imagem(q) for q in selecionadas):
        candidatas_img = [
            q
            for q in pool_valido.to_dict("records")
            if _questao_tem_imagem(q) and str(q["id"]) not in ids_sel
        ]
        if candidatas_img:
            removida = selecionadas[-1]
            nivel_removida = removida.get("dificuldade")
            mesma_dif = [q for q in candidatas_img if q.get("dificuldade") == nivel_removida]
            escolhida = rng.choice(mesma_dif if mesma_dif else candidatas_img)
            selecionadas[-1] = escolhida
            ids_sel.discard(str(removida["id"]))
            ids_sel.add(str(escolhida["id"]))

    return selecionadas


# ── validação ─────────────────────────────────────────────────────────────────

def validar_banco(
    df: pd.DataFrame,
    modelo_nome: str,
    config: dict,
    distribuicao_override: dict[int, float] | None = None,
    regras_assunto: dict[str, dict[str, int]] | None = None,
    aplicar_distribuicao: bool = True,
) -> dict[str, list[str]]:
    """Returns {"erros": [...], "avisos": [...]}:
    - erros: total shortage per subject — block generation.
    - avisos: per-level shortfall — sem compensação automática entre níveis.
    """
    modelo = config[modelo_nome]
    dif_range = modelo.get("dificuldade_range", [1, 5])
    dist_cfg = distribuicao_override or modelo.get("dificuldade_distribuicao")
    regras_norm = _normalizar_regras_assunto(regras_assunto)
    erros: list[str] = []
    avisos: list[str] = []

    for dia_key, dia in modelo["dias"].items():
        for bloco in dia["blocos"]:
            disc = bloco["disciplina"]
            qtd = bloco["quantidade"]

            pool = _filtrar(df, disc, modelo_nome)
            pool = _filtrar_por_assunto(pool, bloco)
            pool_valido = pool[pool["dificuldade"].between(dif_range[0], dif_range[1])]
            regras_disc = regras_norm.get(disc.strip().lower(), {})

            if regras_disc:
                soma_minimos = sum(regras_disc.values())
                if soma_minimos > qtd:
                    erros.append(
                        f"{dia_key} / {disc}: soma dos mínimos por assunto ({soma_minimos}) "
                        f"é maior que a quantidade do bloco ({qtd})."
                    )
                for assunto, minimo in regras_disc.items():
                    disponivel_assunto = int(
                        pool_valido["assunto"]
                        .fillna("")
                        .str.strip()
                        .str.lower()
                        .eq(assunto)
                        .sum()
                    )
                    if disponivel_assunto < minimo:
                        erros.append(
                            f"{dia_key} / {disc} / assunto '{assunto}': necessário {minimo}, "
                            f"disponível {disponivel_assunto}."
                        )

            if len(pool_valido) < qtd:
                erros.append(
                    f"{dia_key} / {disc}: necessário {qtd}, "
                    f"disponível {len(pool_valido)} (dif. {dif_range[0]}–{dif_range[1]})"
                )
                continue

            if aplicar_distribuicao and dist_cfg:
                dist = {int(k): float(v) for k, v in dist_cfg.items()}
                cotas = _calcular_cotas(qtd, dist)
                for nivel, necessario in cotas.items():
                    if necessario == 0:
                        continue
                    disponivel = int((pool_valido["dificuldade"] == nivel).sum())
                    if disponivel < necessario:
                        avisos.append(
                            f"{dia_key} / {disc} / dif.{nivel}: "
                            f"necessário {necessario}, disponível {disponivel} "
                            f"(inclua mais questões desse nível ou reduza a cota desse nível)"
                        )

    return {"erros": erros, "avisos": avisos}


# ── formatação do documento ───────────────────────────────────────────────────

def _limpar_corpo(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _adicionar_cabecalho(doc: Document, nome_banca: str, nome_dia: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nome_banca)
    run.bold = True
    run.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(nome_dia).font.size = Pt(12)

    p3 = doc.add_paragraph(
        "Nome: _____________________________________________   Data: ___/___/______"
    )
    for r in p3.runs:
        r.font.size = Pt(11)
    doc.add_paragraph()


def _adicionar_titulo_bloco(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)


IMG_PATTERN = re.compile(r"^\s*\[IMG\]\s*(\S+)(?:\s*-\s*(.*))?\s*$", re.IGNORECASE)


def _extrair_midia_alternativa(texto_alt: str) -> tuple[str, str] | None:
    match = IMG_PATTERN.match(texto_alt or "")
    if not match:
        return None
    origem = match.group(1).strip()
    legenda = (match.group(2) or "").strip()
    if not origem:
        return None
    return origem, legenda


def _carregar_imagem_bytes(origem: str) -> io.BytesIO:
    parsed = urlparse(origem)
    if parsed.scheme in ("http", "https"):
        with urlopen(origem, timeout=10) as resp:  # nosec B310
            content_type = (resp.headers.get("Content-Type") or "").lower()
            data = resp.read()
        if "svg" in content_type or origem.lower().endswith(".svg"):
            raise ValueError("SVG não é suportado diretamente pelo python-docx.")
        return io.BytesIO(data)
    caminho = Path(origem)
    if not caminho.is_absolute():
        caminho = BASE_DIR / caminho
    if caminho.suffix.lower() == ".svg":
        raise ValueError("SVG não é suportado diretamente pelo python-docx.")
    if not caminho.exists():
        raise FileNotFoundError(f"Imagem não encontrada: {caminho}")
    return io.BytesIO(caminho.read_bytes())


def _adicionar_questao(doc: Document, row: dict, numero: int, letras: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run(f"{numero}. {str(row['texto']).strip()}").font.size = Pt(11)

    for letra in letras:
        texto_alt = str(row.get(letra, "")).strip()
        if texto_alt:
            midia = _extrair_midia_alternativa(texto_alt)
            if midia:
                origem, legenda = midia
                alt = doc.add_paragraph(f"({letra}) {legenda}" if legenda else f"({letra})")
                alt.paragraph_format.left_indent = Cm(1)
                alt.paragraph_format.space_after = Pt(2)
                for r in alt.runs:
                    r.font.size = Pt(11)
                try:
                    img_bytes = _carregar_imagem_bytes(origem)
                    p_img = doc.add_paragraph()
                    p_img.paragraph_format.left_indent = Cm(1.4)
                    run_img = p_img.add_run()
                    run_img.add_picture(img_bytes, width=Cm(6.0))
                except Exception as e:
                    p_fallback = doc.add_paragraph(f"[imagem não renderizada: {origem}]")
                    p_fallback.paragraph_format.left_indent = Cm(1.4)
                    p_fallback.paragraph_format.space_after = Pt(2)
                    for r in p_fallback.runs:
                        r.font.size = Pt(10)
                    p_erro = doc.add_paragraph(f"Motivo: {e}")
                    p_erro.paragraph_format.left_indent = Cm(1.4)
                    p_erro.paragraph_format.space_after = Pt(2)
                    for r in p_erro.runs:
                        r.font.size = Pt(9)
            else:
                alt = doc.add_paragraph(f"({letra}) {texto_alt}")
                alt.paragraph_format.left_indent = Cm(1)
                alt.paragraph_format.space_after = Pt(2)
                for r in alt.runs:
                    r.font.size = Pt(11)

    doc.add_paragraph()


def _adicionar_redacao(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("REDAÇÃO")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    for _ in range(30):
        linha = doc.add_paragraph("_" * 88)
        linha.paragraph_format.space_after = Pt(6)
        for r in linha.runs:
            r.font.size = Pt(10)


# ── interface pública ─────────────────────────────────────────────────────────

def gerar_prova(
    modelo_nome: str,
    caminho_csv: str | Path,
    semente: int | None = None,
    distribuicao_override: dict[int, float] | None = None,
    regras_assunto: dict[str, dict[str, int]] | None = None,
    aplicar_distribuicao: bool = True,
    modo_desenvolvedor: bool = False,
) -> list[Path]:
    """Gera cadernos de prova e gabarito numa pasta timestamped. Retorna arquivos criados."""
    config = carregar_config()
    if modelo_nome not in config:
        raise ValueError(f"Modelo '{modelo_nome}' não encontrado em config_modelos.json")

    df = carregar_banco(caminho_csv)
    modelo = config[modelo_nome]
    dif_range = modelo.get("dificuldade_range", [1, 5])
    distribuicao = (
        {int(k): float(v) for k, v in distribuicao_override.items()}
        if distribuicao_override
        else distribuicao_efetiva(modelo)
    )

    resultado = validar_banco(
        df,
        modelo_nome,
        config,
        distribuicao_override,
        regras_assunto,
        aplicar_distribuicao,
    )
    if resultado["erros"]:
        raise ValueError("Questões insuficientes:\n" + "\n".join(resultado["erros"]))
    if aplicar_distribuicao and resultado["avisos"]:
        raise ValueError(
            "Distribuição por dificuldade inviável sem compensação:\n"
            + "\n".join(resultado["avisos"])
        )

    letras = ["A", "B", "C", "D", "E"][: modelo["alternativas"]]
    redacao_dia = modelo.get("redacao_dia")
    template_path = BASE_DIR / modelo["template"]

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    saida_dir = BASE_DIR / "saida" / f"{modelo_nome}_{timestamp}"
    saida_dir.mkdir(parents=True, exist_ok=True)

    rng = random.Random(semente)
    usadas: set[str] = set()
    gabarito_linhas: list[str] = []
    numero = 1
    arquivos: list[Path] = []

    for dia_key, dia in modelo["dias"].items():
        doc = Document(template_path)
        _limpar_corpo(doc)
        _adicionar_cabecalho(doc, modelo["nome"], dia["nome"])

        for bloco in dia["blocos"]:
            pool = _filtrar(df, bloco["disciplina"], modelo_nome)
            pool = _filtrar_por_assunto(pool, bloco)
            selecionadas = _selecionar_questoes(
                pool,
                bloco["disciplina"],
                bloco["quantidade"],
                dif_range,
                distribuicao,
                rng,
                usadas,
                regras_assunto,
                aplicar_distribuicao,
                forcar_questao_imagem=modo_desenvolvedor,
            )
            _adicionar_titulo_bloco(doc, bloco["exibicao"])
            for q in selecionadas:
                _adicionar_questao(doc, q, numero, letras)
                gabarito_linhas.append(f"{numero:02d} - {str(q['gabarito']).strip()}")
                usadas.add(str(q["id"]))
                numero += 1

        if redacao_dia == dia_key:
            _adicionar_redacao(doc)

        destino = saida_dir / f"{modelo_nome}_{dia_key}.docx"
        doc.save(destino)
        arquivos.append(destino)

    doc_gab = Document()
    p = doc_gab.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"GABARITO — {modelo['nome']}")
    run.bold = True
    run.font.size = Pt(14)
    doc_gab.add_paragraph()
    for linha in gabarito_linhas:
        doc_gab.add_paragraph(linha).paragraph_format.space_after = Pt(1)

    gab_path = saida_dir / f"{modelo_nome}_gabarito.docx"
    doc_gab.save(gab_path)
    arquivos.append(gab_path)

    return arquivos
